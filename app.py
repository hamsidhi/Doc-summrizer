import re
from dataclasses import dataclass
from typing import List, Tuple
import markdown
import mammoth
import streamlit as st
from groq import Groq
from PyPDF2 import PdfReader
import docx
from io import BytesIO
from key_manager import load_api_key, save_api_key


# ==============================
# DATA STRUCTURE
# ==============================
@dataclass
class Document:
    name: str
    extension: str
    text: str
    meta: str


MODEL_FALLBACK_CHAIN = [
    "llama-3.3-70b-versatile",
    "meta-llama/llama-4-maverick-17b-128e-instruct",
    "moonshotai/kimi-k2-instruct-0905",
    "llama-3.1-8b-instant",
]

MODEL_CONTEXT = {
    "llama-3.3-70b-versatile": 128000,
    "meta-llama/llama-4-maverick-17b-128e-instruct": 128000,
    "moonshotai/kimi-k2-instruct-0905": 256000,
    "llama-3.1-8b-instant": 128000,
}


# ==============================
# EXTRACTION
# ==============================
def extract_pdf(file) -> Tuple[str, str]:
    reader = PdfReader(file)
    text = "\n".join([page.extract_text() or "" for page in reader.pages])
    return text.strip(), f"{len(reader.pages)} pages"


def extract_docx(file) -> Tuple[str, str]:
    result = mammoth.extract_raw_text(file)
    return result.value.strip(), "DOCX"


def extract_txt(file) -> Tuple[str, str]:
    return file.getvalue().decode("utf-8", errors="ignore").strip(), "TXT"


def extract_document(file) -> Document:
    ext = file.name.split(".")[-1].lower()

    if ext == "pdf":
        text, meta = extract_pdf(file)
    elif ext == "docx":
        text, meta = extract_docx(file)
    elif ext == "txt":
        text, meta = extract_txt(file)
    else:
        raise ValueError("Unsupported file")

    return Document(file.name, ext, text, meta)


# ==============================
# PROCESSING
# ==============================
def combine_docs(docs: List[Document]) -> str:
    return "\n".join([
        f"========== {d.name} ({d.meta}) ==========\n{d.text}\n"
        for d in docs
    ])


# ==============================
# DETECTION
# ==============================
def detect_flags(text):
    finance = bool(re.search(r"(revenue|profit|cost|\$|%)", text, re.I))
    tech = bool(re.search(r"(api|algorithm|architecture|compliance)", text, re.I))
    ascii_ratio = sum(ord(c) < 128 for c in text[:2000]) / max(len(text[:2000]), 1)
    non_eng = ascii_ratio < 0.9
    return finance, tech, non_eng


# ==============================
# MODEL SELECTION
# ==============================
def select_model(chars, finance, tech, non_eng):
    if finance or tech:
        return "llama-3.3-70b-versatile", "Complex reasoning"

    if non_eng:
        return "meta-llama/llama-4-maverick-17b-128e-instruct", "Multilingual support"

    return "llama-3.3-70b-versatile", "Default"


# ==============================
# CONTEXT SAFETY
# ==============================
# ==============================
# CONTEXT SAFETY — FIXED
# ==============================

# Free tier TPM limits per model (tokens per minute)
MODEL_TPM_LIMIT = {
    "llama-3.3-70b-versatile": 6000,
    "meta-llama/llama-4-maverick-17b-128e-instruct": 6000,
    "moonshotai/kimi-k2-instruct-0905": 6000,
    "llama-3.1-8b-instant": 6000,
}

# Reserve ~1500 tokens for prompt template + output
PROMPT_OVERHEAD = 1500
OUTPUT_RESERVE = 1500

def apply_limit(text, model):
    model_tpm = MODEL_TPM_LIMIT.get(model, 6000)
    
    # Available tokens for document text only
    available_tokens = model_tpm - PROMPT_OVERHEAD - OUTPUT_RESERVE
    available_tokens = max(available_tokens, 500)  # floor safety
    
    # Approx: 1 token ≈ 4 chars
    max_chars = available_tokens * 4  # ≈ 10,000 chars for 6K TPM model

    if len(text) <= max_chars:
        return text, False

    # Smart truncation: take first 60% + last 20% (preserve intro + conclusion)
    head = int(max_chars * 0.6)
    tail = int(max_chars * 0.2)
    truncated = text[:head] + "\n\n[... middle truncated to fit token limit ...]\n\n" + text[-tail:]
    
    return truncated, True

# ==============================
# PROMPT
# ==============================
def build_prompt(text):
    return f"""You are a senior investment analyst and consulting expert. Analyze the document(s) below for a fundraising consultant.

RULES (non-negotiable):
- Every claim → must have a number. No number = don't state it.
- "High/low/good/bad" → always follow with (X% / ₹X / X× benchmark or unit relevant to doc type).
- If a number doesn't exist → say "Not disclosed" — never invent.
- ADAPT every section to the document type. Use relevant units, terminology, and metrics per domain.
- If multiple documents → compare, contrast, and synthesize across all of them.

DOMAIN ADAPTATION GUIDE (auto-apply based on detected type):
- Financial/Startup → Revenue, EBITDA, Burn Rate, Runway, Valuation
- Agricultural → Yield (tons/acre), MSP, Crop cycle, Input cost, Subsidy, Water usage
- Educational → Enrollment, Fee structure, Placement %, Faculty ratio, Accreditation
- Infrastructure/Energy → Capacity (MW/units), Tariff, PLF%, Carbon credits, Capex/Opex
- Healthcare → Patient volume, Cost per bed, Insurance coverage, Regulatory status
- Real Estate → Area (sq ft), Price per sq ft, Occupancy %, ROI timeline
- NGO/Social → Beneficiaries reached, Cost per beneficiary, Grant utilization %
- Other → Extract whatever quantifiable units exist and build analysis around them

---

## 1. DOCUMENT OVERVIEW
- Type: [Financial / Agricultural / Educational / Energy / Healthcare / Real Estate / NGO / Mixed / Other]
- Number of documents: [Single / Multiple — list titles or subjects if multiple]
- Core subject: (1 line)
- Purpose/Goal of the document(s): (1 line)

## 2. KEY METRICS
| Metric | Value | Unit | Investor/Stakeholder Implication |
|--------|-------|------|----------------------------------|
(Extract every quantifiable figure. Adapt metric names to document domain. Min 5 rows.)

## 3. KEY INSIGHTS
(5–7 bullets. Each = Insight + supporting number + why it matters. Domain-appropriate language. No filler.)

## 4. DOMAIN-SPECIFIC ANALYSIS
Primary metrics for this document type:
- [Metric 1 relevant to domain]: ___
- [Metric 2 relevant to domain]: ___
- [Metric 3 relevant to domain]: ___
(Continue for all extractable figures. Mark "Not disclosed" where absent.)

Cost structure:
- Total Input/Operating Cost: ___
- Revenue / Income / Yield Value: ___
- Margin or Surplus: ___ (___%)
- Key cost drivers: (list with values)
- Key income drivers: (list with values)
- Any grants, subsidies, carbon credits, govt schemes: (with values)

## 5. BUSINESS / OPERATING MODEL
(3–5 lines. Explain how value is created and captured — adapted to domain. Embed numbers wherever possible.)

## 6. STRENGTHS
| Strength | Numeric Proof | Domain Context |
|----------|--------------|----------------|
(Min 3. Strength must have a number. Adapt to document type.)

## 7. WEAKNESSES
| Weakness | Estimated Impact | Domain Context |
|----------|-----------------|----------------|
(Internal gaps only. Quantify where possible.)

## 8. RISKS
| Risk | Financial/Operational Impact | Likelihood | Mitigation |
|------|------------------------------|------------|------------|
(Min 3. Risks must be domain-relevant — regulatory, climate, market, execution, etc.)

## 9. OPPORTUNITIES
| Opportunity | Upside Potential | Basis |
|-------------|-----------------|-------|
(Expansion, untapped market, policy tailwinds, tech adoption — quantified.)

## 10. COMPETITIVE / SECTOR POSITION
(2–3 lines. Compare to sector benchmarks, govt targets, or competitors using numbers.
If benchmarks not in document, state "Benchmark not disclosed — qualitative comparison only.")

## 11. MULTI-DOCUMENT SYNTHESIS
(Skip this section if single document.)
- Key agreements across documents:
- Key contradictions or gaps:
- Combined picture for investor/funder: (2–3 lines with numbers)

## 12. FINAL DECISION
- Core value proposition: (1 line with number)
- Risk/reward assessment: (1 line)
- Funding/investment readiness: [Strong / Moderate / Weak] + reason
- Recommendation: PROCEED / CAUTION / PASS
- One-line justification with numbers:

## 13. FINAL SUMMARY
(150–200 words. Written as a briefing note for a senior investor or funder.
Lead with the most important number. End with a clear stance. No filler. Domain-appropriate language.)

---
DOCUMENT(S):
{text}"""


# ==============================
# GROQ CALL
# ==============================
def analyze(api_key, prompt):
    client = Groq(api_key=api_key)
    last_error = None

    for model in MODEL_FALLBACK_CHAIN:
        try:
            st.info(f"Trying model: {model}")

            response = client.chat.completions.create(
                model=model,
                temperature=0.2,
                max_tokens=2048,  # ✅ reduced from 4096 — saves TPM budget
                messages=[
                    {
                        "role": "system",
                        "content": "You are a senior investment analyst. Use numbers. Be precise. No hallucination."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
            )

            st.success(f"✅ Success with: {model}")
            return response.choices[0].message.content.strip()

        except Exception as e:
            last_error = str(e)
            st.warning(f"❌ {model} failed → trying next...")

    raise Exception(f"All models failed. Last error: {last_error}")


# ==============================
# REPORT GENERATION
# ==============================
def create_docx(report_md):
    report_md = report_md.replace('₹', 'Rs.')
    doc = docx.Document()
    doc.add_heading('Document Analysis Report', 0)
    
    lines = report_md.split('\n')
    table_data = []
    in_table = False
    
    for line in lines:
        stripped = line.strip()
        
        if stripped.startswith('|'):
            in_table = True
            cells = [c.strip() for c in stripped.split('|') if c.strip()]
            if cells and not all(c.startswith('-') for c in cells):
                table_data.append(cells)
            continue
        
        if in_table:
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        if j < len(table_data[0]):
                            table.cell(i, j).text = cell
            table_data = []
            in_table = False
            
        if stripped.startswith('###'):
            doc.add_heading(stripped.replace('###', '').strip(), level=3)
        elif stripped.startswith('##'):
            doc.add_heading(stripped.replace('##', '').strip(), level=2)
        elif stripped.startswith('#'):
            doc.add_heading(stripped.replace('#', '').strip(), level=1)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped:
            doc.add_paragraph(stripped)
            
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ==============================
# STREAMLIT UI
# ==============================
def main():
    st.set_page_config(page_title="Document Analyzer", layout="wide")

    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@600&family=Lato:wght@300;400&display=swap');

    html, body, [data-testid="stAppViewContainer"] {
        background-color: #f5f0e8;
        color: #3b2f1e;
        font-family: 'Lato', sans-serif;
    }
    [data-testid="stAppViewContainer"], [data-testid="stHeader"] {
    background-color: #f5f0e8 !important;
    }
    iframe, [data-baseweb="base-input"] {
        background-color: #fdf6ec !important;
        color: #3b2f1e !important;
    }
    [data-testid="stFileUploaderDropzone"] {
    background-color: #fdf6ec !important;
    color: #3b2f1e !important;
    border: none !important;
    }
    [data-testid="stFileUploaderDropzone"] * {
        color: #3b2f1e !important;
        fill: #3b2f1e !important;
    }
    [data-baseweb="file-uploader"] {
        background-color: #fdf6ec !important;
    }
    [data-testid="stSidebar"] {
        background-color: #e8dcc8;
        border-right: 2px solid #c4a97d;
    }
    h1, h2, h3 { font-family: 'Playfair Display', serif; color: #5c3d1e; }
    .stButton > button {
        background-color: #8b5e3c;
        color: #fff8f0;
        border: none;
        border-radius: 6px;
        font-family: 'Lato', sans-serif;
        font-weight: 400;
        padding: 0.5rem 1.2rem;
        transition: background 0.2s;
    }
    .stButton > button:hover { background-color: #6b4226; }
    .stTextInput > div > input {
        background-color: #fdf6ec;
        border: 1px solid #c4a97d;
        color: #3b2f1e;
        border-radius: 6px;
    }
    [data-testid="stFileUploader"] {
        background-color: #fdf6ec;
        border: 2px dashed #c4a97d;
        border-radius: 10px;
        padding: 1rem;
    }
    .stAlert { border-radius: 8px; }
    .stSpinner > div { border-top-color: #8b5e3c !important; }
    .stMarkdown table {
    border-collapse: collapse;
    width: 100%;
    background-color: #fdf6ec;
    border-radius: 8px;
    overflow: hidden;
    margin: 1rem 0;
    }
    th {
        background-color: #8b5e3c;
        color: #fff8f0;
        padding: 10px 14px;
        text-align: left;
        font-family: 'Lato', sans-serif;
    }
    td {
        padding: 9px 14px;
        border-bottom: 1px solid #e8dcc8;
        color: #3b2f1e;
    }
    tr:last-child td { border-bottom: none; }
    tr:hover td { background-color: #f0e6d3; }
    .stDownloadButton > button:hover { background-color: #3b2511; }
    </style>
    """, unsafe_allow_html=True)

    st.markdown("<h1 style='text-align:center;'>📜 Document Intelligence Analyzer</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align:center; color:#8b5e3c; font-size:1.05rem;'>Upload financial, agricultural, educational or any domain document — get investor-grade analysis instantly.</p>", unsafe_allow_html=True)
    st.divider()

    if "key" not in st.session_state:
        st.session_state.key = load_api_key()

    with st.sidebar:
        st.markdown("### 🔐 API Configuration")
        new_key = st.text_input("Groq API Key", value=st.session_state.key, type="password")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("💾 Save"):
                save_api_key(new_key)
                st.session_state.key = new_key
                st.success("Saved")
        with col2:
            if st.button("🗑️ Clear"):
                save_api_key("")
                st.session_state.key = ""
                st.warning("Cleared")

        st.divider()
        st.markdown("**Supported Formats**")
        st.markdown("📄 PDF &nbsp; | &nbsp; 📝 DOCX &nbsp; | &nbsp; 🗒️ TXT", unsafe_allow_html=True)
        st.divider()
        st.markdown("<small style='color:#8b5e3c;'>Free tier: 6K TPM limit.<br>Large docs are smart-truncated.</small>", unsafe_allow_html=True)

    files = st.file_uploader(
        "Drop your documents here",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        label_visibility="visible"
    )

    if files:
        st.markdown(f"<p style='color:#8b5e3c;'>📎 {len(files)} file(s) selected</p>", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    analyze_clicked = st.button("🔍 Analyze Documents", disabled=not files, use_container_width=True)

    if analyze_clicked:

        if not st.session_state.key:
            st.error("API key required. Add it in the sidebar.")
            return

        docs = []
        warnings = []

        with st.spinner("📖 Extracting text from documents..."):
            for f in files:
                try:
                    d = extract_document(f)
                    if d.text:
                        docs.append(d)
                    else:
                        warnings.append(f"{f.name} is empty")
                except Exception as e:
                    warnings.append(f"{f.name} skipped — {str(e)}")

        for w in warnings:
            st.warning(w)

        if not docs:
            st.error("No valid files extracted.")
            return

        combined = combine_docs(docs)
        finance, tech, non_eng = detect_flags(combined)
        model, reason = select_model(len(combined), finance, tech, non_eng)

        st.info(f"🤖 Model: `{model}` | Reason: {reason}")

        safe_text, truncated = apply_limit(combined, model)

        if truncated:
            st.warning(
                f"⚠️ Document truncated: {len(combined):,} → {len(safe_text):,} chars. "
                f"Intro and conclusion preserved."
            )

        with st.spinner("🧠 Analyzing..."):
            report = None
            try:
                prompt = build_prompt(safe_text)
                report = analyze(st.session_state.key, prompt)
            except Exception as e:
                st.error(str(e))
                return

        if report:
            st.success("✅ Analysis complete")
            import markdown
            report_html = markdown.markdown(report, extensions=["tables"])
            st.markdown(report_html, unsafe_allow_html=True)

            # Download Option
            st.divider()
            docx_file = create_docx(report)
            st.download_button(
                label="📥 Download Analysis Report (Word .docx)",
                data=docx_file,
                file_name="analysis_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            st.error("❌ No report generated. Try again.")


       
if __name__ == "__main__":
    main()